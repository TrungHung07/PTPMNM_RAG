from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document


def extract_documents_docx(path: Path) -> list[Document]:
    """
    Bóc tách nội dung từ file DOCX, mỗi đoạn văn (paragraph) thành một Document.

    Lưu ý: Định dạng DOCX không có khái niệm số trang xác định như PDF,
    do đó metadata sử dụng số thứ tự đoạn văn ('paragraph') thay thế.
    Mỗi đoạn văn không rỗng sẽ được đánh index từ 1.

    Args:
        path: Đường dẫn tới file DOCX cần đọc.

    Returns:
        Danh sách Document, mỗi phần tử tương ứng với một đoạn văn có nội dung.
        Metadata gồm: 'paragraph' (số thứ tự đoạn, 1-indexed) và 'source' (tên file).
    """
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    documents: list[Document] = []
    paragraph_index = 0  # Đếm riêng các đoạn CÓ nội dung

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraph_index += 1
            documents.append(Document(
                page_content=text,
                metadata={
                    "paragraph": paragraph_index,   # Số thứ tự đoạn văn có nội dung
                    "source": path.name,            # Tên file dùng để hiển thị citation
                }
            ))
    return documents


def extract_text_docx(path: Path) -> str:
    """
    [Legacy] Bóc tách toàn bộ text từ file DOCX thành một chuỗi duy nhất.
    Hàm này giữ lại để tương thích ngược. Với tính năng citation, dùng
    extract_documents_docx() thay thế.

    Args:
        path: Đường dẫn tới file DOCX cần đọc.

    Returns:
        Toàn bộ nội dung file DOCX dưới dạng một chuỗi văn bản.
    """
    docs = extract_documents_docx(path)
    return "\n".join(doc.page_content for doc in docs)


__all__ = ["extract_text_docx", "extract_documents_docx"]
